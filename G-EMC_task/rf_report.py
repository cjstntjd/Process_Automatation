# -*- coding:utf-8 -*-
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn, nsdecls

# 워드파일이 들어있는 경로를 작성할때, 폴더도 꼭 \\ 이렇게 두번씩 쓰기
#doc = Document('C:\\Users\\84429\\Desktop\\RF\\test.docx')
'''
start_row = 0
for x,paragraph in enumerate(doc.paragraphs):
    if paragraph.text == '3.3 전기적 조건':
        start_row = x
        break
'''
# 시험 모드, 시험 주파수, 시험 환경을 나타내는 표 만들기
def introduction(document):
    style = document.styles['Normal']
    style.font.name='굴림' # 한글 폰트를 따로 설정해 준다
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림') # 한글 폰트를 따로 설정해 준다
    style.font.size = Pt(11)


    document.add_paragraph("") #한 줄 띄기
    table = document.add_table(rows=4, cols=5)
    table.style = document.styles['Table Grid']

    hdr_cells = table.rows[0].cells
    table.rows[0].height = Cm(0.92)
    hdr_cells[0].text = '시 험 모 드'

    a = table.cell(0,1)
    b = table.cell(0,2)
    c = table.cell(0,3)
    d = table.cell(0,4)

    a.merge(b)
    c.merge(d)
    a.merge(c)

    a.text = '2.4G WLAN (802.11b)' # 조작하고 싶은 주파수 대역대 (선택한 id이름으로 )


    for i in range(2):
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # 가운데 정렬
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 가운데 정렬
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[i].height = Cm(0.92)
        hdr_cells[i].width = Cm(3.84)

    hdr_cells = table.rows[1].cells
    table.rows[1].height = Cm(0.92)
    hdr_cells[0].text = '시험 주파수'

    a = table.cell(1,1)
    b = table.cell(1,2)
    c = table.cell(1,3)
    d = table.cell(1,4)

    a.merge(b)
    c.merge(d)
    a.merge(c)

    a.text = 'F1 : 2412MHz	F2 : 2442MHz	F3 : 2472MHz' # F1 , F2, F3 각각 엑셀이랑 연동되도록

    for i in range(2):
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # 가운데 정렬
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 가운데 정렬
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[i].height = Cm(0.92)
        hdr_cells[i].width = Cm(3.84)


    hdr_cells = table.rows[2].cells
    table.rows[2].height = Cm(0.92)
    hdr_cells[0].text = '시 험 환 경'

    a = table.cell(2,0)
    b = table.cell(3,0)

    a.merge(b)

    hdr_cells[1].text = '상    온'
    hdr_cells[2].text = '고    온'
    hdr_cells[3].text = '저    온'
    hdr_cells[4].text = '습    도'

    for i in range(5):
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # 가운데 정렬
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 가운데 정렬
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[i].height = Cm(0.92)
        hdr_cells[i].width = Cm(3.84)

    hdr_cells = table.rows[3].cells # 이것들도 엑셀이랑 연동되도록
    table.rows[3].height = Cm(0.92)
    hdr_cells[1].text = '+25℃'
    hdr_cells[2].text = '+50℃'
    hdr_cells[3].text = '-10℃'
    hdr_cells[4].text = '+35℃, 95%'

    for i in range(5):
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # 가운데 정렬
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 가운데 정렬
        hdr_cells[i].height = Cm(0.92)
        hdr_cells[i].width = Cm(3.84)

    document.add_paragraph("") #한 줄 띄기
    document.add_paragraph("") #한 줄 띄기
    document.add_paragraph("※ 시동 후 1분 경과 이후에 다음의 전기적 조건을 충족시킬 것") #한 줄 띄기

    single_table(document)
    double_table(document)


    document.save('result.docx')


def single_table(d):
    #r : 14 c: 8

    d.add_paragraph("") #한 줄 띄기

    table = d.add_table(rows=14, cols=8)
    table.style = d.styles['Table Grid']

    hdr_cells = table.rows[0].cells
    #table.rows[0].height = Cm(0.92)
    a = table.cell(0,0)
    b = table.cell(0,1)
    c = table.cell(0,2)

    a.merge(b)
    a.merge(c)

    d = table.cell(1,0)
    e = table.cell(1,1)
    f = table.cell(1,2)

    d.merge(e)
    d.merge(f)

    a.merge(d)
    a.text = '시 험 항 목'

    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True
    # height는 다른 열에 따라서 ?
    a.width = Cm(4.83)

    a = table.cell(0,3)
    a.height = Cm(0.89)
    b = table.cell(0,4)
    b.height = Cm(0.89)
    c = table.cell(0,5)
    c.height = Cm(0.89)

    a.merge(b)
    a.merge(c)
    table.rows[0].height = Pt(30)
    a.text = "시  험  결  과"
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True
    a.height = Cm(0.89)
    #a.width = Cm(6.64)
    #a.merge(b)
    #a.merge(c)

    hdr_cells = table.rows[1].cells
    voltage = ['+10% (4.18V)','정격전압 (3.80V)','-#10% (3.42V)']
    for i in range(3,6):
        hdr_cells[i].text = voltage[i-3]
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # 가운데 정렬
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 가운데 정렬
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].height = Pt(30) # 0.89cm
        hdr_cells[i].width = Cm(2.21)

    a = table.cell(0,6)
    b = table.cell(1,6)

    a.merge(b)
    a.text = "합격기준"
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True
    #a.height = Cm(0.89)
    a.width = Cm(3.53)

    a = table.cell(0,7)
    b = table.cell(1,7)

    a.merge(b)
    a.text = "적  부"
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True
    #a.height = Cm(0.89)
    a.width = Cm(1.9)

    a = table.cell(2,0)
    for idb in range(3,14):
        b = table.cell(idb,0)
        #b.height = Pt(20)
        a.merge(b)

    a.text = '주   파   수   허   용   편   차' # 선택한 것들
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    a = table.cell(2,1)
    for idx in range(3,6):
        b = table.cell(idx,1)
        a.merge(b)

    a.text = "F1"
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    temp = ['상 온','고 온','저 온','습 도']
    for t in range(2,6):
        tt = table.cell(t,2)
        tt.width = Cm(2.21)
        tt.text = temp[t-2]
        tt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
        tt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
        #tt.paragraphs[0].runs[0].bold = True

##################### F2

    a = table.cell(6,1)
    for idb in range(7,10):
        b = table.cell(idb,1)
        a.merge(b)

    a.text = "F2"
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    temp = ['상 온','고 온','저 온','습 도']
    for t in range(6,10):
        tt = table.cell(t,2)
        tt.width = Cm(2.21)
        tt.text = temp[t-6]
        tt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
        tt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
        #tt.paragraphs[0].runs[0].bold = True

##################### F3

    a = table.cell(10,1)
    for idb in range(11,14):
        b = table.cell(idb,1)
        a.merge(b)

    a.text = "F3"
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    temp = ['상 온','고 온','저 온','습 도']
    for t in range(10,14):
        tt = table.cell(t,2)
        tt.width = Cm(2.21)
        tt.text = temp[t-10]
        tt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
        tt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
        #tt.paragraphs[0].runs[0].bold = True

    a = table.cell(2,6)
    aa = table.cell(2,7)
    for i in range(3,14):
        b = table.cell(i,6)
        bb = table.cell(i,7)
        a.merge(b)
        aa.merge(bb)


    for rw in range(2,14):
        table.rows[rw].height = Pt(27)

def double_table(d):
    #r : 14 c: 8

    d.add_paragraph("") #한 줄 띄기

    table = d.add_table(rows=26, cols=8)
    table.style = d.styles['Table Grid']

    hdr_cells = table.rows[0].cells
    #table.rows[0].height = Cm(0.92)
    a = table.cell(0,0)
    b = table.cell(0,1)
    c = table.cell(0,2)

    a.merge(b)
    a.merge(c)

    d = table.cell(1,0)
    e = table.cell(1,1)
    f = table.cell(1,2)

    d.merge(e)
    d.merge(f)

    a.merge(d)
    a.text = '시 험 항 목'

    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True
    # height는 다른 열에 따라서 ?
    a.width = Cm(4.83)

    a = table.cell(0,3)
    a.height = Cm(0.89)
    b = table.cell(0,4)
    b.height = Cm(0.89)
    c = table.cell(0,5)
    c.height = Cm(0.89)

    a.merge(b)
    a.merge(c)
    table.rows[0].height = Pt(30)
    a.text = "시  험  결  과"
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True
    a.height = Cm(0.89)
    #a.width = Cm(6.64)
    #a.merge(b)
    #a.merge(c)

    hdr_cells = table.rows[1].cells
    voltage = ['+10% (4.18V)','정격전압 (3.80V)','-#10% (3.42V)']
    for i in range(3,6):
        hdr_cells[i].text = voltage[i-3]
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # 가운데 정렬
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 가운데 정렬
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].height = Pt(30) # 0.89cm
        hdr_cells[i].width = Cm(2.21)

    a = table.cell(0,6)
    b = table.cell(1,6)

    a.merge(b)
    a.text = "합격기준"
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True
    #a.height = Cm(0.89)
    a.width = Cm(3.53)

    a = table.cell(0,7)
    b = table.cell(1,7)

    a.merge(b)
    a.text = "적  부"
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True
    #a.height = Cm(0.89)
    a.width = Cm(1.9)

    a = table.cell(2,0)
    for idb in range(3,14):
        b = table.cell(idb,0)
        #b.height = Pt(20)
        a.merge(b)

    a.text = '안  테  나  공  급  전  력  밀  도' # 선택한 것들
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    a = table.cell(2,1)
    for idx in range(3,6):
        b = table.cell(idx,1)
        a.merge(b)

    a.text = "F1"
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    temp = ['상 온','고 온','저 온','습 도']
    for t in range(2,6):
        tt = table.cell(t,2)
        tt.width = Cm(2.21)
        tt.text = temp[t-2]
        tt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
        tt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
        #tt.paragraphs[0].runs[0].bold = True

##################### F2

    a = table.cell(6,1)
    for idb in range(7,10):
        b = table.cell(idb,1)
        a.merge(b)

    a.text = "F2"
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    temp = ['상 온','고 온','저 온','습 도']
    for t in range(6,10):
        tt = table.cell(t,2)
        tt.width = Cm(2.21)
        tt.text = temp[t-6]
        tt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
        tt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
        #tt.paragraphs[0].runs[0].bold = True

##################### F3

    a = table.cell(10,1)
    for idb in range(11,14):
        b = table.cell(idb,1)
        a.merge(b)

    a.text = "F3"
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    temp = ['상 온','고 온','저 온','습 도']
    for t in range(10,14):
        tt = table.cell(t,2)
        tt.width = Cm(2.21)
        tt.text = temp[t-10]
        tt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
        tt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
        #tt.paragraphs[0].runs[0].bold = True

    a = table.cell(2,6)
    aa = table.cell(2,7)
    for i in range(3,14):
        b = table.cell(i,6)
        bb = table.cell(i,7)
        a.merge(b)
        aa.merge(bb)


############################################################################# no introduction
    a = table.cell(14,0)
    for idb in range(15,26):
        b = table.cell(idb,0)
        #b.height = Pt(20)
        a.merge(b)


    a.text = '점  유  주  파  수  대  역  폭' # 선택한 것들
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    a = table.cell(14,1)
    for idx in range(15,18):
        b = table.cell(idx,1)
        a.merge(b)

    a.text = "F1"
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    temp = ['상 온','고 온','저 온','습 도']
    for t in range(14,18):
        tt = table.cell(t,2)
        tt.width = Cm(2.21)
        tt.text = temp[t-14]
        tt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
        tt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
        #tt.paragraphs[0].runs[0].bold = True

##################### F2

    a = table.cell(18,1)
    for idb in range(19,22):
        b = table.cell(idb,1)
        a.merge(b)

    a.text = "F2"
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    temp = ['상 온','고 온','저 온','습 도']
    for t in range(18,22):
        tt = table.cell(t,2)
        tt.width = Cm(2.21)
        tt.text = temp[t-18]
        tt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
        tt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
        #tt.paragraphs[0].runs[0].bold = True

##################### F3

    a = table.cell(22,1)
    for idb in range(23,26):
        b = table.cell(idb,1)
        a.merge(b)

    a.text = "F3"
    a.width = Cm(1.41)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
    a.paragraphs[0].runs[0].bold = True

    temp = ['상 온','고 온','저 온','습 도']
    for t in range(22,26):
        tt = table.cell(t,2)
        tt.width = Cm(2.21)
        tt.text = temp[t-22]
        tt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #가운데 정렬
        tt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #가운데 정렬
        #tt.paragraphs[0].runs[0].bold = True

    a = table.cell(14,6)
    aa = table.cell(14,7)
    for i in range(15,26):
        b = table.cell(i,6)
        bb = table.cell(i,7)
        a.merge(b)
        aa.merge(bb)


    for rw in range(2,26):
        table.rows[rw].height = Pt(23)



d = Document('C:\\Users\\84430\\Desktop\\FlaskApp - 복사본\\empty.docx')
introduction(d)
